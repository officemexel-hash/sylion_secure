from __future__ import annotations

import html
import json
import re
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "ksiega-4-0-full"
BASELINE_DOCX = ROOT / "SYLION Ksiega v3 4 FIXED.docx"
PHANTOM_DOCX = ROOT / "SYLION_PHANTOM_v3.0.docx"
HTML_OUT = OUT_DIR / "KSIEGA_4_0_FULL_BASELINE_SYLION_PHANTOM.html"
META_OUT = OUT_DIR / "KSIEGA_4_0_FULL_BASELINE_SYLION_PHANTOM.meta.json"


def e(value: object) -> str:
    return html.escape(str(value or ""), quote=True)


def slug(value: str) -> str:
    value = value.lower()
    value = re.sub(r"[^a-z0-9ąćęłńóśźż]+", "-", value)
    return value.strip("-")[:80] or "section"


def iter_block_items(parent: DocxDocument) -> Iterable[Paragraph | Table]:
    body = parent.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def table_to_html(table: Table) -> str:
    rows = []
    for r_idx, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            text = "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
            tag = "th" if r_idx == 0 else "td"
            cells.append(f"<{tag}>{e(text)}</{tag}>")
        rows.append("<tr>" + "".join(cells) + "</tr>")
    return '<table class="docx-table">' + "".join(rows) + "</table>"


def paragraph_to_html(paragraph: Paragraph) -> str:
    text = paragraph.text.strip()
    if not text:
        return ""
    style = paragraph.style.name if paragraph.style else ""
    escaped = e(text)
    if style.startswith("Heading 1"):
        return f'<h1 id="{slug(text)}">{escaped}</h1>'
    if style.startswith("Heading 2"):
        return f'<h2 id="{slug(text)}">{escaped}</h2>'
    if style.startswith("Heading 3"):
        return f'<h3 id="{slug(text)}">{escaped}</h3>'
    if style.startswith("Heading 4"):
        return f'<h4 id="{slug(text)}">{escaped}</h4>'
    if text.startswith(("• ", "- ")):
        return f"<p class=\"bullet\">{escaped}</p>"
    if len(text) < 120 and text.isupper():
        return f"<h3>{escaped}</h3>"
    return f"<p>{escaped}</p>"


def extract_docx_html(path: Path, *, title: str, max_blocks: int | None = None) -> tuple[str, dict]:
    doc = Document(str(path))
    blocks = []
    paragraphs = 0
    tables = 0
    headings = 0
    for block in iter_block_items(doc):
        if max_blocks and len(blocks) >= max_blocks:
            break
        if isinstance(block, Paragraph):
            rendered = paragraph_to_html(block)
            if rendered:
                paragraphs += 1
                if rendered.startswith(("<h1", "<h2", "<h3", "<h4")):
                    headings += 1
                blocks.append(rendered)
        elif isinstance(block, Table):
            tables += 1
            blocks.append(table_to_html(block))
    meta = {
        "source": str(path.name),
        "paragraphs_rendered": paragraphs,
        "tables_rendered": tables,
        "headings_rendered": headings,
    }
    return f'<section class="source-doc"><h1>{e(title)}</h1>{"".join(blocks)}</section>', meta


def svg(nodes, edges, title: str, width: int = 1060, height: int = 520) -> str:
    parts = [
        f'<figure class="svg-figure"><figcaption>{e(title)}</figcaption>',
        f'<svg viewBox="0 0 {width} {height}" role="img" aria-label="{e(title)}" xmlns="http://www.w3.org/2000/svg">',
        '<defs><marker id="arrow" markerWidth="10" markerHeight="10" refX="8" refY="3" orient="auto" markerUnits="strokeWidth"><path d="M0,0 L0,6 L9,3 z" fill="#255f85"/></marker></defs>',
        f'<rect x="0" y="0" width="{width}" height="{height}" rx="18" fill="#f7fbff" stroke="#b7c9d8"/>',
        f'<text x="24" y="36" font-size="22" font-family="Segoe UI, Arial" font-weight="700" fill="#0a2342">{e(title)}</text>',
    ]
    node_map = {n["id"]: n for n in nodes}
    for a, b, label in edges:
        na = node_map[a]
        nb = node_map[b]
        x1 = na["x"] + na["w"]
        y1 = na["y"] + na["h"] / 2
        x2 = nb["x"]
        y2 = nb["y"] + nb["h"] / 2
        if nb["x"] < na["x"]:
            x1 = na["x"] + na["w"] / 2
            y1 = na["y"] + na["h"]
            x2 = nb["x"] + nb["w"] / 2
            y2 = nb["y"]
        parts.append(f'<line x1="{x1}" y1="{y1}" x2="{x2}" y2="{y2}" stroke="#255f85" stroke-width="2.2" marker-end="url(#arrow)"/>')
        if label:
            parts.append(f'<text x="{(x1+x2)/2}" y="{(y1+y2)/2-6}" font-size="12" font-family="Segoe UI, Arial" fill="#255f85">{e(label)}</text>')
    for n in nodes:
        fill = n.get("fill", "#ffffff")
        stroke = n.get("stroke", "#2f80ed")
        parts.append(f'<rect x="{n["x"]}" y="{n["y"]}" width="{n["w"]}" height="{n["h"]}" rx="10" fill="{fill}" stroke="{stroke}" stroke-width="2"/>')
        label = n["label"]
        lines = re.split(r"\n| / ", label)
        for i, line in enumerate(lines[:4]):
            parts.append(f'<text x="{n["x"]+14}" y="{n["y"]+26+i*18}" font-size="{14 if i else 15}" font-family="Segoe UI, Arial" font-weight="{700 if i == 0 else 500}" fill="#102a43">{e(line[:42])}</text>')
    parts.append("</svg></figure>")
    return "".join(parts)


def diagrams_html() -> str:
    diagrams = []
    diagrams.append(svg(
        [
            {"id": "portal", "label": "Portal publiczny\nzakup tokenu", "x": 30, "y": 90, "w": 180, "h": 70},
            {"id": "admin", "label": "Admin API\nprivate control plane", "x": 270, "y": 90, "w": 190, "h": 70},
            {"id": "op", "label": "Operator record\nsubskrypcja i pakiety", "x": 520, "y": 90, "w": 200, "h": 70},
            {"id": "g1", "label": "G1\nnetwork gateway", "x": 90, "y": 240, "w": 160, "h": 70},
            {"id": "g2", "label": "G2\naccess broker", "x": 330, "y": 240, "w": 160, "h": 70},
            {"id": "wl", "label": "WORKLOAD\nFirecracker/KVM", "x": 570, "y": 240, "w": 190, "h": 70},
            {"id": "cdr", "label": "CDR + Audit\nmetadata only", "x": 810, "y": 240, "w": 190, "h": 70},
            {"id": "term", "label": "Pixel/laptop\nthin client", "x": 30, "y": 390, "w": 190, "h": 70},
            {"id": "router", "label": "Puli AX\naccess router", "x": 280, "y": 390, "w": 190, "h": 70},
            {"id": "apps", "label": "Apps\nSignal/Telegram/etc.", "x": 590, "y": 390, "w": 220, "h": 70},
        ],
        [
            ("portal", "admin", "claim"),
            ("admin", "op", "bootstrap"),
            ("op", "g1", "provision"),
            ("g1", "g2", "tunnel"),
            ("g2", "wl", "broker"),
            ("wl", "cdr", "events"),
            ("term", "router", "Wi-Fi"),
            ("router", "g1", "VPN"),
            ("wl", "apps", "microVM"),
            ("apps", "term", "pixels"),
        ],
        "D1. Architektura logiczna SYLION 4.0",
        height=500,
    ))
    diagrams.append(svg(
        [
            {"id": "px", "label": "Pixel GrapheneOS\nterminal posture", "x": 25, "y": 120, "w": 190, "h": 80},
            {"id": "pa", "label": "Puli AX\nkill switch / VPN", "x": 255, "y": 120, "w": 180, "h": 80},
            {"id": "g1", "label": "G1\nfirst trust gate", "x": 475, "y": 120, "w": 160, "h": 80},
            {"id": "g2", "label": "G2\nsession broker", "x": 675, "y": 120, "w": 160, "h": 80},
            {"id": "fc", "label": "Firecracker\napp microVM", "x": 875, "y": 120, "w": 160, "h": 80},
            {"id": "sig", "label": "Signal / Telegram\nWhatsApp / Threema\nZangi / Office", "x": 730, "y": 310, "w": 260, "h": 95},
            {"id": "aud", "label": "Monitoring\nCDR / audit / SIEM", "x": 370, "y": 310, "w": 250, "h": 95},
            {"id": "adm", "label": "Admin panel\nblue-team view", "x": 70, "y": 310, "w": 220, "h": 95},
        ],
        [
            ("px", "pa", "device pair"),
            ("pa", "g1", "VPN 1"),
            ("g1", "g2", "VPN 2"),
            ("g2", "fc", "session"),
            ("fc", "sig", "runtime"),
            ("fc", "aud", "metadata"),
            ("g1", "aud", "telemetry"),
            ("aud", "adm", "alerts"),
        ],
        "D2. Ścieżka Pixel - Puli AX - G1 - G2 - workload",
    ))
    diagrams.append(svg(
        [
            {"id": "pay", "label": "Payment\nStripe/CoinGate/Mollie", "x": 40, "y": 100, "w": 220, "h": 75},
            {"id": "tok", "label": "Token service\nhash + scope", "x": 320, "y": 100, "w": 200, "h": 75},
            {"id": "claim", "label": "Claim\none-time action", "x": 580, "y": 100, "w": 190, "h": 75},
            {"id": "boot", "label": "Bootstrap\noperator record", "x": 830, "y": 100, "w": 190, "h": 75},
            {"id": "pkg", "label": "Packages\nPixel + Puli AX", "x": 180, "y": 300, "w": 210, "h": 75},
            {"id": "infra", "label": "Infrastructure\nG1/G2/workload", "x": 470, "y": 300, "w": 220, "h": 75},
            {"id": "audit", "label": "Audit\npayment/token/provision", "x": 760, "y": 300, "w": 220, "h": 75},
        ],
        [
            ("pay", "tok", "webhook"),
            ("tok", "claim", "token"),
            ("claim", "boot", "create"),
            ("boot", "pkg", "download"),
            ("boot", "infra", "provision"),
            ("infra", "audit", "events"),
            ("tok", "audit", "ledger"),
        ],
        "D3. Portal, płatność, token i bootstrap operatora",
    ))
    diagrams.append(svg(
        [
            {"id": "pool", "label": "Provider pool\ncountries/capabilities", "x": 35, "y": 90, "w": 230, "h": 80},
            {"id": "policy", "label": "Tier policy\nlimits + rotation", "x": 320, "y": 90, "w": 210, "h": 80},
            {"id": "alloc", "label": "Allocator\ncost + isolation", "x": 590, "y": 90, "w": 210, "h": 80},
            {"id": "new", "label": "New VPS\nhigh tier", "x": 850, "y": 90, "w": 170, "h": 80},
            {"id": "reuse", "label": "Sanitized reuse\nlower tier", "x": 850, "y": 250, "w": 170, "h": 80},
            {"id": "wipe", "label": "Wipe/reinstall\nkey destruction", "x": 590, "y": 250, "w": 210, "h": 80},
            {"id": "audit", "label": "Evidence\ncleanup + attestation", "x": 320, "y": 250, "w": 210, "h": 80},
        ],
        [
            ("pool", "policy", "eligible"),
            ("policy", "alloc", "decision"),
            ("alloc", "new", "dedicated"),
            ("new", "wipe", "release"),
            ("wipe", "audit", "prove"),
            ("audit", "reuse", "return to pool"),
            ("alloc", "reuse", "standard/pro"),
        ],
        "D4. Rotacja jurysdykcyjna i reuse VPS po operatorze",
    ))
    diagrams.append(svg(
        [
            {"id": "upload", "label": "Ingress\nfile/input", "x": 40, "y": 120, "w": 170, "h": 70},
            {"id": "class", "label": "Classification\ntype/risk/hash", "x": 260, "y": 120, "w": 190, "h": 70},
            {"id": "cdr", "label": "CDR\nsanitize/disarm", "x": 500, "y": 120, "w": 180, "h": 70},
            {"id": "policy", "label": "Policy\nallow/deny/quarantine", "x": 730, "y": 120, "w": 230, "h": 70},
            {"id": "app", "label": "App microVM\nsafe handoff", "x": 730, "y": 300, "w": 230, "h": 70},
            {"id": "audit", "label": "Audit\nmetadata only", "x": 410, "y": 300, "w": 220, "h": 70},
            {"id": "block", "label": "Quarantine\noperator notice", "x": 110, "y": 300, "w": 220, "h": 70},
        ],
        [
            ("upload", "class", "inspect"),
            ("class", "cdr", "transform"),
            ("cdr", "policy", "decision"),
            ("policy", "app", "allow"),
            ("policy", "block", "deny"),
            ("cdr", "audit", "evidence"),
        ],
        "D5. CDR i kontrola przepływu plików",
    ))
    diagrams.append(svg(
        [
            {"id": "req", "label": "Requirement\nexpected behavior", "x": 40, "y": 90, "w": 220, "h": 70},
            {"id": "test", "label": "Human test\nPixel/ADB/laptop", "x": 320, "y": 90, "w": 220, "h": 70},
            {"id": "evi", "label": "Evidence\nscreens/logs/metadata", "x": 600, "y": 90, "w": 220, "h": 70},
            {"id": "res", "label": "PASS/FAIL\nBLOCKED/UNKNOWN", "x": 870, "y": 90, "w": 160, "h": 70},
            {"id": "def", "label": "Defect\nsmallest module", "x": 600, "y": 280, "w": 220, "h": 70},
            {"id": "fix", "label": "Repair\nminimal patch", "x": 320, "y": 280, "w": 220, "h": 70},
            {"id": "retest", "label": "Retest\nsame test", "x": 40, "y": 280, "w": 220, "h": 70},
        ],
        [
            ("req", "test", "run"),
            ("test", "evi", "capture"),
            ("evi", "res", "classify"),
            ("res", "def", "if fail"),
            ("def", "fix", "patch"),
            ("fix", "retest", "repeat"),
            ("retest", "test", "same path"),
        ],
        "D6. Antyhalucynacyjny cykl test - naprawa - retest",
    ))
    diagrams.append(svg(
        [
            {"id": "super", "label": "Global superadmin\napps/providers/tiers", "x": 45, "y": 80, "w": 230, "h": 80},
            {"id": "admin", "label": "Admin\noperators/billing/audit", "x": 330, "y": 80, "w": 230, "h": 80},
            {"id": "ops", "label": "SRE/SOC\nmonitoring/incidents", "x": 615, "y": 80, "w": 230, "h": 80},
            {"id": "op", "label": "Operator\nown environments only", "x": 330, "y": 275, "w": 230, "h": 80},
            {"id": "portal", "label": "Portal user\ntoken claim only", "x": 615, "y": 275, "w": 230, "h": 80},
        ],
        [
            ("super", "admin", "delegates"),
            ("admin", "ops", "alerts"),
            ("admin", "op", "provisions"),
            ("portal", "admin", "token event"),
            ("ops", "admin", "incident"),
        ],
        "D7. Role i uprawnienia paneli",
    ))
    diagrams.append(svg(
        [
            {"id": "pixel", "label": "Pixel\nGrapheneOS", "x": 90, "y": 90, "w": 210, "h": 80},
            {"id": "router", "label": "Puli AX\nrouter package", "x": 430, "y": 90, "w": 210, "h": 80},
            {"id": "fido", "label": "FIDO2\nuser presence", "x": 770, "y": 90, "w": 210, "h": 80},
            {"id": "gate", "label": "PHANTOM admission\nall three required", "x": 345, "y": 285, "w": 360, "h": 95, "fill": "#eef8f4", "stroke": "#25a18e"},
        ],
        [
            ("pixel", "gate", "posture"),
            ("router", "gate", "network"),
            ("fido", "gate", "touch"),
        ],
        "D8. PHANTOM: trójca sprzętowa",
    ))
    return '<section class="diagrams"><h1>Rozdział D. Diagramy i grafy Księgi 4.0</h1>' + "".join(diagrams) + "</section>"


def front_matter() -> str:
    return f"""
    <section class="cover">
      <h1>Księga 4.0</h1>
      <h2>SYLION Secure + PHANTOM v3.0</h2>
      <p class="subtitle">Pełny baseline techniczny, architektura, modułowość, panele, tiering, rotacje, analiza zagrożeń, PHANTOM governance i aktualny stan wdrożenia.</p>
      <table class="meta">
        <tr><th>Data</th><td>{date.today().isoformat()}</td></tr>
        <tr><th>Źródła</th><td>SYLION Księga v3.4 FIXED, PHANTOM v3.0, dokumenty admin-panel-v2, implementacja services/admin-api.</td></tr>
        <tr><th>Status</th><td>Baseline 4.0 draft do dalszej redakcji i human gate. Nie jest claimem pełnej produkcyjności.</td></tr>
        <tr><th>Granica bezpieczeństwa</th><td>Opis defensywnej architektury i governance. Brak sekretów, brak executorów RF/telecom, brak instrukcji obchodzenia zabezpieczeń.</td></tr>
      </table>
    </section>
    <section>
      <h1>Rozdział 0. Korekta względem poprzedniego artefaktu</h1>
      <p>Poprzedni PDF był krótkim skrótem. Księga 4.0 w tej wersji ma inną funkcję: jest dokumentem bazowym i śladem scalającym pełną Księgę 3.4, profil PHANTOM v3.0 i aktualny stan implementacji. Z tego powodu zawiera pełny import materiału źródłowego Księgi 3.4, osobną warstwę aktualizacji 4.0, diagramy oraz profil PHANTOM przepisany do formy bezpiecznej architektonicznie.</p>
      <p>Najważniejsza zasada interpretacyjna: jeżeli tekst odziedziczony z Księgi 3.4 albo PHANTOM v3.0 jest sprzeczny z warstwą aktualizacji 4.0, nadrzędna jest warstwa 4.0. Szczególnie dotyczy to routera Puli AX, granic PHANTOM, portalu zakupowego, tokenów, nowego tieringu i polityki reuse VPS.</p>
    </section>
    <section>
      <h1>Rozdział 1. Normatywny baseline Księgi 4.0</h1>
      <table>
        <tr><th>Obszar</th><th>Wymaganie</th><th>Status</th></tr>
        <tr><td>Terminal</td><td>Terminal jest cienkim klientem i nie przechowuje danych roboczych operatora.</td><td>Baseline</td></tr>
        <tr><td>G1/G2</td><td>Każdy operator ma indywidualne G1 i G2. Współdzielenie tych warstw wymaga ADR i human gate.</td><td>Baseline</td></tr>
        <tr><td>Workload</td><td>Workloady są izolowane per operator i aplikacja zgodnie z tierem. Pro i wyżej wymagają Firecracker albo silniejszej izolacji.</td><td>Baseline/PHANTOM</td></tr>
        <tr><td>CDR</td><td>CDR jest obowiązkowe dla plików i przepływów między strefami.</td><td>Baseline</td></tr>
        <tr><td>Portal</td><td>Portal publiczny działa na osobnym VPS i komunikuje się z prywatnym Admin API wyłącznie przez allowlistowane endpointy.</td><td>Baseline</td></tr>
        <tr><td>Providerzy</td><td>Administrator zarządza provider registry: kraje, koszty, KVM, Firecracker, TDX, SEV-SNP, bare metal.</td><td>Baseline</td></tr>
        <tr><td>Rotacja</td><td>Rotacja jest funkcją subskrypcji i policy engine, nie dowolnym przyciskiem operatora.</td><td>Baseline</td></tr>
        <tr><td>PHANTOM</td><td>PHANTOM jest profilem rozszerzonym, nie może automatycznie odblokowywać baseline execution.</td><td>PHANTOM</td></tr>
        <tr><td>RF/telecom</td><td>Elementy RF, SIM, modem i identyfikatory są lab-only governance, bez executorów produktu.</td><td>Lab-only</td></tr>
      </table>
    </section>
    <section>
      <h1>Rozdział 2. Nowe elementy 4.0 względem Księgi 3.4</h1>
      <ul>
        <li>Puli AX jest aktualnym routerem roboczym dla ścieżki mobilnej. Starsze odniesienia do innych routerów są legacy albo wymagają ponownej kwalifikacji.</li>
        <li>Portal zakupowy jest oddzielną strefą publiczną, z tokenami, płatnościami Stripe/CoinGate/Mollie i scenariuszem resellerów.</li>
        <li>Minimalny publiczny okres subskrypcji wynosi 12 miesięcy. Starsze 6-miesięczne guardy administracyjne wymagają harmonizacji.</li>
        <li>Każdy operator ma indywidualne G1 i G2. Workload bare metal może być współdzielony tylko w niższych tierach i tylko po izolacji oraz sanitizacji.</li>
        <li>PHANTOM i Sovereign wymagają dedykowanych albo operator-only workloadów.</li>
        <li>Reuse VPS po innym operatorze jest dopuszczalne tylko po wipe, reinstall, zniszczeniu sekretów, rekwalifikacji i dowodzie audytowym.</li>
        <li>Panel operatora musi umożliwiać kontrolę własnych środowisk: reset, recreate, aplikacje, session TTL, backup, panic, Matrix, rotacje i monitoring własnej ścieżki.</li>
        <li>Panel administratora musi być konsolą blue-team: operatorzy, providerzy, koszty, alerty, anomalie, CDR, rotacje, status G1/G2/workload.</li>
        <li>Testy muszą być factual: aplikacja działa dopiero wtedy, gdy człowiek/ADB potwierdzi używalność, a nie tylko HTTP 200 albo działający proces.</li>
      </ul>
    </section>
    """


def generated_reference_sections() -> str:
    tiers = [
        ("Pilot", "99 EUR", "1 188 EUR", "6", "shared pool / kontenery", "brak domyślnie"),
        ("Standard", "199 EUR", "2 388 EUR", "10", "shared pool / kontenery", "manualna ograniczona"),
        ("Pro", "499 EUR", "5 988 EUR", "20", "Firecracker pool", "harmonogram i provider rotation"),
        ("Phantom", "1 000 EUR", "12 000 EUR", "40", "dedykowany / operator-only", "pełna polityka"),
        ("Sovereign", "2 999 EUR", "35 988 EUR", "60", "operator-only dedicated", "pełna polityka + najwyższa izolacja"),
    ]
    tier_rows = "".join(f"<tr><td>{e(a)}</td><td>{e(b)}</td><td>{e(c)}</td><td>{e(d)}</td><td>{e(f)}</td><td>{e(g)}</td></tr>" for a, b, c, d, f, g in tiers)
    admin_rows = [
        ("Operatorzy", "tworzenie, status, tier, koszt, G1/G2/workload, pakiety, audit"),
        ("Providerzy", "kraje, regiony, KVM, Firecracker, TDX, SEV-SNP, bare metal, ceny"),
        ("Subskrypcje", "ceny, terminy, limity, upgrade, tokeny, resellerzy"),
        ("Monitoring", "alerty, anomalie, zmiany kluczy, błędy CDR, status tuneli"),
        ("PHANTOM", "hardening gates, lab-only records, human gate, brak automatycznego execution"),
        ("CDR", "polityki plików, quarantine, hash, decyzje allow/deny"),
    ]
    operator_rows = [
        ("Sesja", "licznik czasu, odblokowanie, hasła G1/G2/workload, FIDO2 później"),
        ("Aplikacje", "Signal, Telegram, WhatsApp, Threema, Zangi, DuckDuckGo, LibreOffice, Exodus"),
        ("Środowiska", "liczba instancji według tieru, reset, recreate, prepare new session"),
        ("Rotacje", "kraje, providerzy, częstotliwość i polityki tylko w granicach tieru"),
        ("Backup/Panic", "backup operatora, auto-wipe po nieaktywności, panic code z poziomami skutku"),
        ("Matrix", "wniosek o własny serwer Matrix jako opcjonalny addon"),
    ]
    return f"""
    <section>
      <h1>Rozdział 3. Tiering, ceny i limity 4.0</h1>
      <table><tr><th>Tier</th><th>Cena/mies.</th><th>Cena roczna</th><th>Środowiska</th><th>Workload</th><th>Rotacja</th></tr>{tier_rows}</table>
      <p>CDR jest obowiązkowe we wszystkich tierach. G1 i G2 są indywidualne dla każdego operatora. Różnica między tierami dotyczy przede wszystkim liczby środowisk, rodzaju izolacji workloadów, prawa do rotacji, dostępu do Matrix, dedykacji hostów oraz wymagań PHANTOM.</p>
    </section>
    <section>
      <h1>Rozdział 4. Funkcjonalności panelu administratora</h1>
      <table><tr><th>Zakładka / obszar</th><th>Funkcjonalność</th></tr>{''.join(f'<tr><td>{e(a)}</td><td>{e(b)}</td></tr>' for a,b in admin_rows)}</table>
      <p>Panel administratora jest prywatną konsolą sterowania i blue-team. Nie powinien być wystawiony publicznie. Musi pokazywać koszt utrzymania każdego operatora, jego tier, aktualny status G1/G2/workload, rotacje, alerty i problemy produkcyjne.</p>
    </section>
    <section>
      <h1>Rozdział 5. Funkcjonalności panelu operatora</h1>
      <table><tr><th>Zakładka / obszar</th><th>Funkcjonalność</th></tr>{''.join(f'<tr><td>{e(a)}</td><td>{e(b)}</td></tr>' for a,b in operator_rows)}</table>
      <p>Operator nie zarządza globalną platformą. Zarządza własnym środowiskiem: sesją, aplikacjami, resetem microVM/kontenerów, backupiem, panic, Matrix i rotacją dostępną w jego tierze.</p>
    </section>
    <section>
      <h1>Rozdział 6. Polityka reuse VPS i rotacji po operatorze</h1>
      <p>Niższe tiery mogą być przenoszone na oczyszczone zasoby po innych operatorach. Warunkiem jest pełny cleanup: zakończenie sesji, odpięcie operatora, zniszczenie sekretów, wipe danych, reinstall albo rebuild, nowe certyfikaty, nowe klucze, nowy binding, dowód audytowy i brak otwartych alertów.</p>
      <p>Wysokie tiery, zwłaszcza Phantom i Sovereign, wymagają dedykacji albo formalnej rekwalifikacji zasobu. To oznacza, że zasób po innym operatorze może zostać użyty tylko po uznaniu go za czysty, pusty i zgodny z profilem dedykowanym.</p>
      <table>
        <tr><th>Tier</th><th>Reuse G1/G2</th><th>Reuse workload</th><th>Warunek</th></tr>
        <tr><td>Pilot</td><td>nie dla aktywnych operatorów</td><td>tak, z puli</td><td>wipe + nowe sekrety + audit</td></tr>
        <tr><td>Standard</td><td>nie dla aktywnych operatorów</td><td>tak, z puli</td><td>wipe + reinstall + audit</td></tr>
        <tr><td>Pro</td><td>nie dla aktywnych operatorów</td><td>tak, Firecracker pool</td><td>cleanup gates + capacity + provider policy</td></tr>
        <tr><td>Phantom</td><td>tylko po pełnej rekwalifikacji</td><td>dedykowany/operator-only</td><td>human gate + attestation + compliance approval</td></tr>
        <tr><td>Sovereign</td><td>zasadniczo brak współdzielenia</td><td>operator-only</td><td>formalna rekwalifikacja albo nowy host</td></tr>
      </table>
    </section>
    """


def phantom_safe_profile() -> str:
    return """
    <section class="phantom">
      <h1>CZĘŚĆ PHANTOM 4.0 — profil rozszerzony</h1>
      <p>PHANTOM 4.0 jest profilem wysokiego ryzyka. Nie jest pojedynczym modułem, lecz zestawem wymagań dla terminala, routera, G1, G2, workloadów, streamingu, CDR, rotacji, audytu i procedur. PHANTOM nie może automatycznie odblokowywać funkcji baseline bez human gate.</p>
      <h2>PHANTOM: warstwy</h2>
      <table>
        <tr><th>Warstwa</th><th>Cel</th><th>Mechanizm</th><th>Ryzyko rezydualne</th></tr>
        <tr><td>Pixel</td><td>cienki terminal</td><td>GrapheneOS, profil CA, posture check, brak danych lokalnych</td><td>aktywny malware może widzieć bieżący ekran/wejście</td></tr>
        <tr><td>Puli AX</td><td>kontrolowany WAN</td><td>VPN, kill switch, DNS leak prevention, pakiet routera</td><td>baseband/router firmware/supply chain</td></tr>
        <tr><td>FIDO2</td><td>fizyczna obecność operatora</td><td>touch/biometria, odnowienie sesji</td><td>procedury odzyskiwania i social engineering</td></tr>
        <tr><td>G1</td><td>pierwsza brama</td><td>certyfikaty, VPN, posture gate</td><td>metadane i DoS</td></tr>
        <tr><td>G2</td><td>broker sesji</td><td>limit sesji, routing, docelowo blind broker</td><td>jawny stream jeśli broker nie jest E2EE</td></tr>
        <tr><td>Workload</td><td>izolacja aplikacji</td><td>Firecracker, dedykacja, rebuild</td><td>VM escape, kernel hosta, provider</td></tr>
        <tr><td>CDR</td><td>kontrola plików</td><td>sanitize, quarantine, hash, audit</td><td>formaty nieobsłużone i błędy polityki</td></tr>
        <tr><td>Blue team</td><td>wykrywanie naruszeń</td><td>metadane, anomalia, alerty, eBPF roadmap</td><td>brak treści oznacza ograniczenia diagnostyczne</td></tr>
      </table>
      <h2>PHANTOM: czego nie wolno twierdzić</h2>
      <ul>
        <li>Nie wolno twierdzić, że system daje niewykrywalność albo anonimowość absolutną.</li>
        <li>Nie wolno twierdzić, że RF fingerprinting da się w pełni rozwiązać software'owo.</li>
        <li>Nie wolno twierdzić, że VPN chroni metadane operatora komórkowego.</li>
        <li>Nie wolno twierdzić, że G2 jest ślepy, dopóki streaming nie ma E2EE od mikroVM do terminala.</li>
        <li>Nie wolno uruchamiać funkcji RF/telecom jako zwykłych funkcji produktu.</li>
      </ul>
      <h2>PHANTOM: wymagania wdrożeniowe</h2>
      <table>
        <tr><th>Wymaganie</th><th>Test</th><th>Gate</th></tr>
        <tr><td>Pixel widzi G1 tylko po przejściu posture policy</td><td>ADB human regression + negative test bez routera/FIDO2</td><td>CISO/Architect</td></tr>
        <tr><td>Puli AX wymusza tunel i kill switch</td><td>test DNS leak, route leak, reconnect, WAN loss</td><td>SRE/Security</td></tr>
        <tr><td>G2 nie zapisuje streamu ani wejścia</td><td>log audit + storage negative test</td><td>Security</td></tr>
        <tr><td>Workload recreate niszczy stan aplikacji</td><td>human test: zamknij, recreate, uruchom od zera</td><td>Platform</td></tr>
        <tr><td>CDR działa dla plików</td><td>test upload/download/quarantine/sanitize</td><td>Security</td></tr>
        <tr><td>Rotacja nie zostawia sekretów</td><td>cleanup evidence + new cert fingerprint + audit</td><td>SRE</td></tr>
      </table>
    </section>
    """


def css() -> str:
    return """
    @page { size: A4; margin: 14mm 13mm 16mm; }
    body { font-family: "Segoe UI", Arial, sans-serif; color: #172033; font-size: 9.5pt; line-height: 1.42; }
    h1 { color:#082340; font-size:22pt; line-height:1.08; margin: 22px 0 10px; page-break-before: always; border-bottom:2px solid #2f80ed; padding-bottom:7px; }
    .cover h1 { page-break-before: auto; font-size:42pt; border:0; margin-top:70px; }
    h2 { color:#0a3558; font-size:14pt; margin: 16px 0 8px; page-break-after: avoid; }
    h3 { color:#123f63; font-size:11.5pt; margin: 13px 0 6px; page-break-after: avoid; }
    h4 { color:#123f63; font-size:10.5pt; margin: 10px 0 5px; }
    p { margin: 6px 0; }
    .subtitle { font-size: 14pt; max-width: 760px; color:#314963; }
    table { width:100%; border-collapse: collapse; margin: 8px 0 14px; page-break-inside: avoid; font-size:8pt; }
    th, td { border:1px solid #c6d2dd; padding:5px 6px; vertical-align:top; }
    th { background:#eef5fb; color:#082340; text-align:left; }
    .meta th { width: 24%; }
    .bullet { margin-left: 18px; }
    .source-doc { page-break-before: always; }
    .svg-figure { page-break-inside: avoid; margin: 12px 0 18px; }
    .svg-figure figcaption { font-weight: 700; color:#0a3558; margin-bottom: 6px; }
    svg { width: 100%; height: auto; display:block; }
    .diagrams h1, .phantom h1 { page-break-before: always; }
    .notice { border-left: 5px solid #2f80ed; background:#f2f7fc; padding:10px 12px; margin:12px 0; }
    """


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    baseline_html, baseline_meta = extract_docx_html(BASELINE_DOCX, title="CZĘŚĆ B — Import pełnej Księgi SYLION v3.4 FIXED jako baseline traceability")
    phantom_source_doc = Document(str(PHANTOM_DOCX))
    phantom_headings = []
    for p in phantom_source_doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style = p.style.name if p.style else ""
        if "Heading" in style:
            phantom_headings.append((style, text))
    phantom_index_rows = "".join(f"<tr><td>{e(style)}</td><td>{e(text)}</td></tr>" for style, text in phantom_headings)
    phantom_index = f"""
    <section>
      <h1>CZĘŚĆ C — Indeks źródłowy PHANTOM v3.0</h1>
      <p>Poniższy indeks pokazuje strukturę źródłowego dokumentu PHANTOM v3.0. Księga 4.0 nie importuje automatycznie operacyjnych procedur wysokiego ryzyka. Zamiast tego przepisuje PHANTOM jako kontrolowany profil architektoniczny, governance i threat model.</p>
      <table><tr><th>Styl</th><th>Nagłówek źródłowy</th></tr>{phantom_index_rows}</table>
    </section>
    """
    body = "\n".join([
        front_matter(),
        diagrams_html(),
        generated_reference_sections(),
        baseline_html,
        phantom_safe_profile(),
        phantom_index,
    ])
    html_doc = f"""<!doctype html>
    <html lang="pl">
    <head>
      <meta charset="utf-8">
      <title>Księga 4.0 Full Baseline SYLION PHANTOM</title>
      <style>{css()}</style>
    </head>
    <body>{body}</body>
    </html>"""
    HTML_OUT.write_text(html_doc, encoding="utf-8")
    META_OUT.write_text(json.dumps({
        "html": str(HTML_OUT),
        "baseline": baseline_meta,
        "phantom_headings": len(phantom_headings),
        "generated": date.today().isoformat(),
        "safety": "PHANTOM operational telecom/RF executor details are not imported as product instructions.",
    }, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"html": str(HTML_OUT), "meta": str(META_OUT), "baseline": baseline_meta, "phantom_headings": len(phantom_headings)}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
